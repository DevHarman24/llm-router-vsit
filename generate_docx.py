import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doubts = [
    # Category 1: JEE Preparation & Strategy (2026)
    ("I wasted my 11th grade. Is it still possible to crack JEE 2026 if I start now, or should I focus only on boards?", "https://www.reddit.com/r/JEENEETards/"),
    ("Should I take a drop for JEE 2026 or join a tier-3 private engineering college this year?", "https://www.reddit.com/r/JEENEETards/"),
    ("Are the new reduced NCERT textbooks enough for JEE Mains 2026 Chemistry, or do I need older editions?", "https://www.reddit.com/r/JEENEETards/"),
    ("There is a mismatch between my marked answers and the recorded response sheet for JEE Main 2026. How do I challenge it?", "https://www.reddit.com/r/JEENEETards/"),
    ("Which is better for JEE 2026: Offline coaching or online platforms like PW/Unacademy?", "https://www.quora.com/search?q=JEE+2026"),
    ("The mathematics section in the 2026 January shift was extremely lengthy. How can I improve my speed and accuracy for the April attempt?", "https://www.reddit.com/r/JEENEETards/"),
    ("Is HC Verma still relevant for JEE Advanced 2026 Physics, or should I directly jump to Irodov?", "https://www.quora.com/search?q=JEE+Advanced+2026"),
    ("What documents are required for the EWS certificate verification during JEE 2026 JoSAA counselling?", "https://www.reddit.com/r/JEENEETards/"),

    # Category 2: NEET Preparation & Strategy (2026)
    ("I scored 575 in NEET 2026. Should I take a drop year, or is it better to look for MBBS abroad?", "https://www.reddit.com/r/JEENEETards/"),
    ("What are the controversial bonus questions in the NEET 2026 Physics paper? Will NTA accept multiple options?", "https://www.reddit.com/r/JEENEETards/"),
    ("Are state quota cutoffs expected to increase significantly for NEET 2026 compared to last year?", "https://www.quora.com/search?q=NEET+2026+cutoff"),
    ("Is biology alone enough to score 360/360 in NEET 2026, or do I need out-of-syllabus reference books?", "https://www.quora.com/search?q=NEET+2026+Biology"),
    ("Physics is my weakest subject. How should I approach numericals for NEET 2026 when I struggle with basic math?", "https://www.reddit.com/r/IndianAcademia/"),
    ("Is Allen's test series tougher than the actual NEET 2026 exam level?", "https://www.quora.com/search?q=NEET+2026+Allen"),
    ("Can I change my exam center for NEET 2026 after submitting the final application form?", "https://www.reddit.com/r/JEENEETards/"),

    # Category 3: CUET & College Admissions (2026)
    ("Will CUET 2026 be merged with JEE and NEET, or will it remain a completely separate exam?", "https://www.reddit.com/r/IndianAcademia/"),
    ("What is the normalization process for CUET 2026, and how will it affect Delhi University (DU) cutoffs?", "https://www.reddit.com/r/IndianAcademia/"),
    ("If I want to pursue B.Sc. Physics at St. Stephen's, which specific domain subjects must I choose for CUET 2026?", "https://www.reddit.com/r/cuetards/"),
    ("Are private universities like Ashoka, Krea, and Flame worth the high fee compared to top CUET colleges?", "https://www.reddit.com/r/IndianAcademia/"),
    ("Is there any negative marking in the General Test section for CUET 2026?", "https://www.quora.com/search?q=CUET+2026"),

    # Category 4: CBSE Boards & Academic Structure (2026)
    ("Will CBSE strictly conduct board exams twice a year starting in 2026 as per the new NEP guidelines?", "https://www.reddit.com/r/CBSE/"),
    ("Does my 10th and 12th board percentage matter for placements if I get into a top NIT/IIT in 2026?", "https://www.reddit.com/r/IndianAcademia/"),
    ("What are the passing criteria for CBSE 2026 if I fail in one core subject but pass in my additional subject?", "https://www.reddit.com/r/CBSE/"),
    ("Is taking a 'dummy school' for 11th and 12th still a safe and viable option for JEE/NEET 2026 aspirants?", "https://www.reddit.com/r/JEENEETards/"),
    ("How do I handle the pressure and burnout of preparing for pre-boards and competitive exams simultaneously?", "https://www.reddit.com/r/CBSE/"),

    # Category 5: General & International Admissions (2026)
    ("Is it too late to start building extracurriculars in 11th grade for US College admissions (Class of 2026)?", "https://www.reddit.com/r/ApplyingToCollege/"),
    ("What is the procedure for a fee refund if I withdraw my seat after the 2nd round of counselling in 2026?", "https://www.reddit.com/r/IndianAcademia/"),
    ("Are online degrees (like IIT Madras BS in Data Science) considered valid for masters abroad in 2026?", "https://www.reddit.com/r/IndianAcademia/"),
    ("I am scoring well but feeling extremely burnt out and depressed. Should I take a break from my 2026 prep?", "https://www.reddit.com/r/JEENEETards/")
]

doc = docx.Document()
doc.add_heading('2026 Student FAQs and Doubts', 0)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Query / Doubt'
hdr_cells[1].text = 'Source'

for query, source in doubts:
    row_cells = table.add_row().cells
    row_cells[0].text = query
    row_cells[1].text = source

doc.save('Student_Doubts_2026.docx')
print("Successfully created Student_Doubts_2026.docx")
