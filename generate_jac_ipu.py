import urllib.request
import json
import docx
import re
import ssl

ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE

links = [
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGW0eOV0IT1k6OFBH9YrSNXE9_tTX7LLY22hJkrriTu2eoOtF5lmW1yzpivVvBHrT5yFaI_XyARkpH45fmCfZlpQ4ly4WE72X0TQYhcW0gNaUKPUJm0exOb_3uplJ4qbZj1kFBXB3MItCwlNp0dkFTWFReyLz2D7jaE-_ptanyAvlQ=",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQETZQMeMA8J0829lVaLR1MBK3sOrGDH6YP7OibB2bKrvucHOUI8kWRZlo3BZe0uIZu3pB0rbJ12q9xgRYESiXss4xscIMaoSBuWcqfOqTmgHY4gNbWxaM2hd0hNoEpoYBF0qWd7-wc55RSOOkTjlmIFpWHTPcydblKuzqY1eIqI6h1f2AhKQLKE",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGgszgVUhVOtNS1whdwWkTxPQgQEkGvNKxfOKfe--1pEJsr8RILDOD1SNiXEgBKWEZhKfc4ananMv0lzP6VFYXWFxle_C3VU5vZgiSpM0e8hw3vTTOT2hbnsgPMPjMEg-KCLL3C7BmUir3NbK_e4mrdDqDyf6QOliMlieDt0-744g==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQFUSiLeAkpvCuHHyGW_MA-hEqhQgcpDYsxpXOM_s9ca_R80X7i_gD34gLJfVBydW2qpmRNw7NANm7VRAdAxIXRX5RL6JFmVtNPmyh9m3MP1X2pPs_a6sIrpnkoH0DMTBY3Iny4-ojNBshV2HJDatjHSnMPidMbyycp9xGdCNuuCwqR3",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQG0HjsZcdoXggcFBG14j7vUQMAOlEW8LGtwi76jn7_4hCjVN2Fer6FOd9gAgC0dLvtpr7vQQzbXuijGwmzfQZupYmOEwO4MwiDRrCUKuS5ekfp4uiQaUHZlSHfxITsIKjG0-ZfB9YEyq7Bj0gpPkDtwIjJqAIxq8ic8FSY9qM7b4ORkkFdzN4Ek7GpKzCqsz084znRis9XBZLc=",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQFijlrgHN7xp6CCGHOyNVJ-XaQuYevRGc2k71JT0DMwTRtUwQ52ui28JTAbzZ1YftZxX_YZ_XNGnCIhMsKUgPGiDSzj_qpZxW1G1rkJdD6TfsCLqNlcK10Tc8ldRl2IPNkyoLMStS9hpdCp_Ue3z1Db4YRfo2SkLAlEHlnRU6oWFSINfNcgStPdY7rrbke-kipHwZ-xNM5Gk6mn_zo=",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQEFXmpROuUVEZuceMHwFcez4vmzifB1ucyEL2f07BwmQnknMbsx5LDnUoamj-xkZOphcVq0eHrkuspscJmJcAYu3-R2N7FUc4Szq3JkAypRqBRDxWzV3vpsffeJX2tM-WtywEOdj4-qN7NaK9aZmCR5qniKYmxBduslE1zAZDmP0OIRKqylJfJBe1fep0JlBV0=",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQHozwZhjbkpMd5_s2Uxs1oTv4jdEKxuW_LwdIMthPgTz7tuyXHd5SepI3KianhVF5whLTzb0AeeeJE8HqCRUfKHqDOA5j7tdzPNkYbDXzTlb-m65wEfhwGOpBtHTGXnlUoRwgCyRO6eTsxCkvX33Ra4GSyABi5pevBoLzuWyBoGtOYHto-pi7P35UeyInc5DItUMGVbguHl4P4vow==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQHCtRn2w74yPcQdYKEN9RgXQt8KcME9rj135Qxz9Y6Map1geCRMdWYuYXBd9d_q-aWcRS78qZHr_-esYBUp5A5m6c04J_kVZ9rAhW3BFFbVmgGY4ocWMUd_XKSXe6gFofg6BRxQQFc3R433OahoXeLOa0eIR7HMWe9YDNSy7KPGTbN3YUly4_R12OziwMX_6qAd",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQHjMN8JOcAmU56V7FQODK2GXkrLyEPUX0H7PWT6GPrdCmTcSPhSvmBqMZUqurxzfY-ybQHz2mvhvgXQdKObpLoQ5yuEDVj9ZQt5p19mW9lnxIKG0pcfiSj-wykYlAeCsnrKvJl8sEW5qaqPdS-Z1T42U1i0DbcuFtvTKIlswHRwTw==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQG5UK5rk-BvlHV8yKFyQ-Fx9kbDe2Cx8nYJqCr_TE0KUXFZQ8koRwHME-vjjxG1C-_ZDFsq0BsWFuQT68qHNSBoe1TIi_rUZt-HACq8IPBLQunbu2ojwBhJ5Tsdl3C2IMqsC1HGEo_VworxjlqHM5X-s68gV3JwGkyCFhCI20-dNZ2yZz157ilbcTFFLmAS7vY=",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGmpgx3oPuplIdqxQ4B5UXpuo9tg5GjepN8pZu91hevFAPfzCLdMdH8GRzmA9wSrBcC6Ym-vOCOyJBp3Ox0cOJwvDpOmw-SYjLWt_7jvKlxnCZK8CRvRgtiENkH9tGTaoo4Q0wJZQdN4iE9kEQ3I_ELwDfDLmoAOes9b-NuIUjoE6QSRgM9CCrdROev5F9YKBYrR3Ag-Yk3qkmm",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQEHZz1H8H3Z_MLWJCWkSNkG2aYqECSM8NwHE7ZSTj5oYhkgVDJXM1vQnEp8Tp0l6ygIQhi7YngrxUGxEfxF_o-jNIi566DoTAIzyhAWQxoUUvVwrq-66-Q6Ic6EdOKlh_aJlfRLJ54S66SkMQ8jWNFtAq5j0_hIYn5iYaXcOUh3e7HZhiloEfkbeABuxAXJNEcBlzuP4FtTJQ==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGyou10ntoLYVahUh2h2XFER9E-RwUN7ju19gVTW1dXWroiNRW8BLHDM3SMfxtBc7CJSRPK1emgCrp769PFhESGHhtaX-zQcfHVMRyh8hMXBqDpq5-8uZPwQ9OJ2zXdheYW3cfhqjzLlnYb2opf3-g6-9fzO7iWwGZTyMLxAXGw",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGpzo_DvRdjrOlZk1_AOUFRH-iTMpRXzHFqGReXOYPdAlGlpvk2kPm_VJZBrIXKGDZ9SJIv7a8FQyzPJOmrpwtUbYbT7UhcfRRa4WLwcCgh__pTM73mjOyroW58yeSneyUOlbIsTGlTPdVogV7TVjGKD1f3W-cc8KLuK5Qt_t2GKgM=",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGLz5u1yIyu6QjtSBbfgs-Rxg0uZElYH0bYTkEX6jOeMuviDMSBHBDCxQiHhk_HfUK6Pkqbvid45gIO-6knUyI0R_ie5sXpVbMd0RAOltWMF9WZsCLTDVziCEzK1oyCbQDJHzKc3c1YlnWs2LVkJHZzFwxci2VxBk9XTZZe_wY2kDS-7-FaeDQx1s2BJjRBBMKS1f1XbR5BgtZ8sw==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQFiowvbn5yshTfy737Q07j1VRD63evsfk1DF0y7kJup1YHmsMaPNwaSV2UzHsB0TCoh82yHHmfXESKeW1cnPftiBwfiUUBwntoorbHdO2G0EHk1uP35aoz4NzxVckZntn7oJ-1eq1daLu8sDXto5oYg-PTcfbtVVyqtGcisPUo=",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGU-ABOgXAQQYPP9mYDZwtGGoO1CyipoW33JPUSabww_vvukD4idK7p1qKMczCcnyNIFavDQOiOJP4ZkHWiFwb8CLdqlNPYcBEvVuRmqg_wadOwR8B3O5hZK1QnLcODqF0hqcnV3w99_MaR2I2JnaVPsbwCSnb8fqTaYNHsBdm2fUTgzALIwB0_",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQHKfekJzTeizflnJAKhHnOWrNslHtJNA5kncIBOMYKmeU9gk2sowOaRLO0FD6JZkHhJQZWwu7ywfJUlpU-RGzremNx2jHHGOVcCaiqy9zzsRl2QG76SWirnnbya-11Mn8LIkWa6ELgiAxoZQCSo5ot81kIGyHwMyNeyMyna4Yxt97_GpYscD3RoZZMINA==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQHDg8mB9rYIAIrQPHWVRm3kHiaJ_98l6gBjbzU-r61DzaaI6wA6i7FYYoqiipr-TX2oBWCqLJOcS39_n79M-FXzQj8MvTMSHtllBihntWdyhMA4wuIREXJXZvcx1GhvgnqfpQza3PbluaGvUP74SwxGYsxRa5Ki3ICRBT5E6hIRXgnOOakBJjCku7FAMFcWbKSJ-PJY",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQEpauPqLP3wjpYPej7c9G_WzHrrshqA4Gcy-9j_qRW8JjR74aSaZsMPNBYeYhZJNeu6p9MN55qLZbJJy8gjizB8rhBq0RASheuJJK9xTEchGUV9NRVl5C5wcejngaISnN4xpa0Eks3EkA5lNB6WmtBmWJ19XL2VoTvWSIm7FB7Fv-5yVO-OZOg7UQ==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQGN_pKGdJb-yX_FK2fhVnSAzH8qBmg5OjpPL0zP8ybDlnbU_j6mFs4jKPLA6L6kW9-rhTiojqo-RZ0YiuTMggNc3jcKKloUduwjUnZ24SaBXKpETcLiZX5m9c6HnI3aYMF1TGWVOQ9YjMHFO7ts2beTT-Fs0VztB23NLKCBchehdG1fUG9ZYuo8H_qcrFBrBkrnJ01pOg==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQF1zO89T6IiHh_ZkyumMC2QFwHBcMr1bry8mumbFMz2vBU587SjlIvNv6bZaymIs_3FAgSl65lHOdC82Bqmr_yB5XtkktxZsI4WUT-3B-hwQuQz69i-pGGjzh8VPml8_tKseJPgS6DmGsiTAUpVkAI-MIvUhEicst_jLTO695M8OA==",
    "https://vertexaisearch.cloud.google.com/grounding-api-redirect/AUZIYQFKa-u4Kq0eOzIhzxsNVaHNqTCXOcaIdJPa3J0VpfvUjxKRN6QWA1BZQlxzmstby-o3qAJbCxST5HxxFWvpak1k8Z6_Zu9SYX_hVeibtDZLue9FQpQt5wT-NONuhj84szO7I8eF10Rax6W-D8i6EirMN4o2AizkWOluz4mcKT_Z_Yp-C4s="
]

def format_title(slug):
    title = slug.replace('_', ' ').strip()
    if title:
        title = title[0].upper() + title[1:] + "?"
    return title

unique_doubts = {}

print("Resolving links...")
for url in links:
    try:
        req = urllib.request.Request(url, method='HEAD')
        opener = urllib.request.build_opener(urllib.request.HTTPRedirectHandler)
        res = opener.open(req, timeout=5)
        final_url = res.url
        # extract slug
        # usually: https://www.reddit.com/r/Btechtards/comments/13wcmlh/jac_counselling_doubt/
        parts = final_url.split('/comments/')
        if len(parts) > 1:
            slug = parts[1].split('/')[1]
            title = format_title(slug)
            if title and final_url not in unique_doubts.values():
                unique_doubts[title] = final_url
    except Exception as e:
        print(f"Error resolving {url}: {e}")

print(f"Found {len(unique_doubts)} unique doubts.")

# If we don't have enough, let's append a few manually crafted highly relevant ones with simulated accurate urls for demonstration
# But since we have 24 links, hopefully we get at least 15-20 unique ones.

doc = docx.Document()
doc.add_heading('JAC Delhi & IPU Counselling Doubts', 0)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Query / Doubt'
hdr_cells[1].text = 'Exact Source Link'

for query, source in unique_doubts.items():
    row_cells = table.add_row().cells
    row_cells[0].text = query
    row_cells[1].text = source

doc.save('JAC_IPU_Counselling_Doubts.docx')
print("Successfully created JAC_IPU_Counselling_Doubts.docx")
